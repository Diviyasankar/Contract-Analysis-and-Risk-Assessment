"""
Document Loader Module
Handles extraction of text from PDF, DOCX, and TXT files
"""

import os
import re
from typing import Optional, Tuple
from pathlib import Path

# PDF Processing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX Processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentLoader:
    """
    Unified document loader for PDF, DOCX, and TXT files
    """
    
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.txt']
    
    def __init__(self):
        self.text = ""
        self.metadata = {}
        self.language = "en"
    
    def load(self, file_path: str) -> Tuple[str, dict]:
        """
        Load and extract text from a document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {extension}")
        
        self.metadata = {
            "filename": path.name,
            "extension": extension,
            "size_bytes": path.stat().st_size if path.exists() else 0
        }
        
        if extension == '.pdf':
            self.text = self._extract_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            self.text = self._extract_docx(file_path)
        elif extension == '.txt':
            self.text = self._extract_txt(file_path)
        
        # Detect language
        self.language = self._detect_language(self.text)
        self.metadata["language"] = self.language
        self.metadata["char_count"] = len(self.text)
        self.metadata["word_count"] = len(self.text.split())
        
        return self.text, self.metadata
    
    def load_from_bytes(self, file_bytes: bytes, filename: str) -> Tuple[str, dict]:
        """
        Load document from bytes (for Streamlit file uploader)
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        import tempfile
        
        extension = Path(filename).suffix.lower()
        
        # Write to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            return self.load(tmp_path)
        finally:
            # Cleanup temp file
            os.unlink(tmp_path)
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_parts = []
        
        # Try pdfplumber first (better for complex PDFs)
        if PDF_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    
                    self.metadata["page_count"] = len(pdf.pages)
                
                if text_parts:
                    return "\n\n".join(text_parts)
            except Exception as e:
                print(f"pdfplumber failed: {e}, trying PyPDF2...")
        
        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                self.metadata["page_count"] = len(reader.pages)
                return "\n\n".join(text_parts)
            except Exception as e:
                raise RuntimeError(f"Failed to extract PDF text: {e}")
        
        raise RuntimeError("No PDF library available. Install pdfplumber or PyPDF2.")
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            self.metadata["paragraph_count"] = len(paragraphs)
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX text: {e}")
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise RuntimeError("Failed to read text file with any known encoding")
    
    def _detect_language(self, text: str) -> str:
        """
        Detect if text is primarily English or Hindi
        
        Returns:
            'en' for English, 'hi' for Hindi, 'mixed' for bilingual
        """
        if not text:
            return "en"
        
        # Hindi Unicode range: \u0900-\u097F (Devanagari)
        hindi_pattern = re.compile(r'[\u0900-\u097F]')
        hindi_chars = len(hindi_pattern.findall(text))
        
        # English/ASCII pattern
        english_pattern = re.compile(r'[a-zA-Z]')
        english_chars = len(english_pattern.findall(text))
        
        total = hindi_chars + english_chars
        if total == 0:
            return "en"
        
        hindi_ratio = hindi_chars / total
        
        if hindi_ratio > 0.7:
            return "hi"
        elif hindi_ratio > 0.2:
            return "mixed"
        else:
            return "en"
    
    def get_sections(self, text: str) -> list:
        """
        Split document into logical sections based on headings
        
        Returns:
            List of (heading, content) tuples
        """
        # Common section patterns in contracts
        section_patterns = [
            r'^(?:ARTICLE|Article|SECTION|Section|CLAUSE|Clause)\s*[\d\.]+[:\.]?\s*(.+)$',
            r'^\d+\.\s+([A-Z][A-Za-z\s]+)$',
            r'^[IVXLC]+\.\s+(.+)$',
            r'^(?:SCHEDULE|Schedule|ANNEXURE|Annexure|EXHIBIT|Exhibit)\s*[\w\-]*[:\.]?\s*(.*)$'
        ]
        
        combined_pattern = '|'.join(f'({p})' for p in section_patterns)
        
        lines = text.split('\n')
        sections = []
        current_heading = "Introduction"
        current_content = []
        
        for line in lines:
            is_heading = False
            for pattern in section_patterns:
                if re.match(pattern, line.strip(), re.IGNORECASE):
                    # Save previous section
                    if current_content:
                        sections.append((current_heading, '\n'.join(current_content)))
                    current_heading = line.strip()
                    current_content = []
                    is_heading = True
                    break
            
            if not is_heading:
                current_content.append(line)
        
        # Add last section
        if current_content:
            sections.append((current_heading, '\n'.join(current_content)))
        
        return sections


# Quick test
if __name__ == "__main__":
    loader = DocumentLoader()
    print("DocumentLoader initialized successfully!")
    print(f"PDF support: {PDF_AVAILABLE}")
    print(f"DOCX support: {DOCX_AVAILABLE}")
